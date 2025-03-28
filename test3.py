import streamlit as st
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
import time
import pytesseract
from PIL import Image
import cv2
import numpy as np
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import os
import re
import shutil
from tqdm import tqdm

stop_extraction = False  # Global flag to stop extraction

def setup_driver():
    options = webdriver.ChromeOptions()
    driver = webdriver.Chrome(options=options)
    return driver

def preprocess_image(image_path):
    img = cv2.imread(image_path)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    gray = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
    return gray

def clean_text(text):
    text = re.sub(r"[^a-zA-Z0-9.,!?()\-:; ]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text

def extract_text_from_canvas(driver, slide_number):
    img_path = f"slides_images/slide_{slide_number}.png"
    driver.save_screenshot(img_path)
    processed_img = preprocess_image(img_path)
    pil_img = Image.fromarray(processed_img)
    raw_text = pytesseract.image_to_string(pil_img)
    extracted_text = clean_text(raw_text)
    return extracted_text, img_path

def extract_slides(prezi_url, total_slides):
    global stop_extraction
    st.write("Starting extraction...")
    driver = setup_driver()
    driver.get(prezi_url)
    time.sleep(5)
    os.makedirs("slides_images", exist_ok=True)
    doc = Document()
    doc.add_heading("Extracted Slides Data", level=1)
    progress_bar = st.progress(0)
    
    for i in range(total_slides):
        if stop_extraction:
            st.warning("⏹️ Extraction Stopped!")
            break
        extracted_text, img_path = extract_text_from_canvas(driver, i+1)
        doc.add_heading(f"Slide {i+1}", level=2)
        doc.add_paragraph(extracted_text)
        doc.add_picture(img_path, width=Inches(5.5))
        st.image(img_path, caption=f"Slide {i+1}")
        st.write(f"**Extracted Text:** {extracted_text}")
        body = driver.find_element(By.TAG_NAME, "body")
        body.send_keys(Keys.ARROW_RIGHT)
        time.sleep(3)
        progress_bar.progress((i + 1) / total_slides)
    
    word_file = "Extracted_Presentation.docx"
    doc.save(word_file)
    driver.quit()
    convert(word_file)
    pdf_path = "Extracted_Presentation.pdf"
    return word_file, pdf_path

def cleanup_files():
    shutil.rmtree("slides_images", ignore_errors=True)
    os.remove("Extracted_Presentation.docx")
    os.remove("Extracted_Presentation.pdf")

st.title("📑 Prezi Slide Extractor")
prezi_url = st.text_input("Enter Prezi URL:")
total_slides = st.number_input("Number of slides:", min_value=1, max_value=100, value=10)
user_email = st.text_input("Enter email to receive files (Optional):")

if st.button("Start Extraction"):
    if prezi_url:
        stop_extraction = False
        word_file, pdf_file = extract_slides(prezi_url, total_slides)
        st.success("✅ Extraction Complete!")
        st.download_button("Download Word File", open(word_file, "rb"), file_name=word_file)
        st.download_button("Download PDF File", open(pdf_file, "rb"), file_name=pdf_file)
        cleanup_files()

if st.button("Stop Extraction"):
    stop_extraction = True
